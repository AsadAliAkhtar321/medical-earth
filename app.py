from flask import Flask, render_template, request, send_file, jsonify, session
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from PIL import Image
import tempfile
import base64
import weasyprint
import markdown
from ebooklib import epub
import uuid
import shutil
import io
from PyPDF2 import PdfMerger
import cv2
import zipfile
import threading
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'  # Change this in production
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

class SubtitleProcessor:
    def __init__(self, subtitles_per_page=3, no_spacing=True, narrow_borders=True, add_bookmarks=True,
                 heading_style="numbered", dark_theme=False, text_color="#000000", label_color="#FF0000",
                 page_organization=True, create_folder=True, custom_title="Video Subtitle Report"):
        self.subtitles_per_page = subtitles_per_page
        self.no_spacing = no_spacing
        self.narrow_borders = narrow_borders
        self.add_bookmarks = add_bookmarks
        self.heading_style = heading_style
        self.dark_theme = dark_theme
        self.text_color = text_color
        self.label_color = label_color
        self.page_organization = page_organization
        self.create_folder = create_folder
        self.custom_title = custom_title
        self.is_processing = False

    def log_message(self, message, log_list):
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_list.append(f"[{timestamp}] {message}")

    def parse_srt_time(self, time_str):
        time_str = time_str.replace(',', '.')
        h, m, s = time_str.split(':')
        return int(h) * 3600 + int(m) * 60 + float(s)

    def format_time(self, seconds):
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        return f"{hours:02d}:{minutes:02d}:{secs:06.3f}"

    def parse_srt_file(self, srt_path, log_list):
        subtitles = []
        try:
            with open(srt_path, 'r', encoding='utf-8') as file:
                content = file.read()
            blocks = re.split(r'\n\s*\n', content.strip())
            for block in blocks:
                lines = block.strip().split('\n')
                if len(lines) >= 3:
                    subtitle_num = lines[0].strip()
                    timing_line = lines[1].strip()
                    start_time, end_time = timing_line.split(' --> ')
                    text = '\n'.join(lines[2:])
                    subtitles.append({
                        'number': subtitle_num,
                        'start_time': self.parse_srt_time(start_time),
                        'end_time': self.parse_srt_time(end_time),
                        'text': text
                    })
            self.log_message(f"Parsed {len(subtitles)} subtitles from {srt_path}", log_list)
        except Exception as e:
            self.log_message(f"Error parsing SRT file {srt_path}: {str(e)}", log_list)
        return subtitles

    def capture_screenshot(self, video_path, timestamp, log_list):
        cap = cv2.VideoCapture(video_path)
        cap.set(cv2.CAP_PROP_POS_MSEC, timestamp * 1000)
        ret, frame = cap.read()
        cap.release()
        if ret:
            screenshot = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            self.log_message(f"Screenshot captured for timestamp {self.format_time(timestamp)}", log_list)
            return screenshot
        self.log_message(f"Failed to capture screenshot for timestamp {self.format_time(timestamp)}", log_list)
        return None

    def get_heading_text(self, subtitle, index):
        if self.heading_style == "numbered":
            return f"Subtitle {subtitle['number']}"
        elif self.heading_style == "time":
            start = self.format_time(subtitle['start_time'])
            end = self.format_time(subtitle['end_time'])
            return f"Scene {start} - {end}"
        else:
            return f"Scene {index + 1}"

    def hex_to_rgb(self, hex_color):
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    def create_output_folder(self, output_path, relative_path, base_output_dir):
        if not self.create_folder:
            return output_path, None
        output_dir = os.path.dirname(output_path)
        folder_name = f"{os.path.splitext(os.path.basename(output_path))[0]}_output"
        folder_path = os.path.join(base_output_dir, relative_path, folder_name) if relative_path else os.path.join(base_output_dir, folder_name)
        os.makedirs(folder_path, exist_ok=True)
        images_folder = os.path.join(folder_path, "images")
        os.makedirs(images_folder, exist_ok=True)
        new_output_path = os.path.join(folder_path, os.path.basename(output_path))
        return new_output_path, images_folder

    def create_docx_report(self, subtitles, screenshots, output_path, relative_path, video_name, base_output_dir, log_list):
        output_path, images_folder = self.create_output_folder(output_path, relative_path, base_output_dir)
        doc = Document()
        if self.narrow_borders:
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
        if self.dark_theme:
            section = doc.sections[0]
            sectPr = section._sectPr
            pgBg = OxmlElement("w:background")
            pgBg.set(qn("w:color"), "1a1a1a")
            sectPr.append(pgBg)
        else:
            section = doc.sections[0]
            sectPr = section._sectPr
            pgBg = OxmlElement("w:background")
            pgBg.set(qn("w:color"), "ffffff")
            sectPr.append(pgBg)
        title = doc.add_heading(self.custom_title or 'Video Subtitle Report', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(0)
        for run in title.runs:
            run.font.color.rgb = RGBColor(*self.hex_to_rgb(self.text_color)) if not self.dark_theme else RGBColor(255, 255, 255)
        summary_p = doc.add_paragraph()
        summary_p.paragraph_format.space_before = Pt(0)
        summary_p.paragraph_format.space_after = Pt(0)
        summary_p.paragraph_format.line_spacing = 1.0
        summary_run = summary_p.add_run('Total subtitles: ')
        summary_run.bold = True
        summary_run.font.color.rgb = RGBColor(*self.hex_to_rgb(self.text_color)) if not self.dark_theme else RGBColor(255, 255, 255)
        count_run = summary_p.add_run(str(len(subtitles)))
        count_run.font.color.rgb = RGBColor(*self.hex_to_rgb(self.text_color)) if not self.dark_theme else RGBColor(255, 255, 255)
        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_before = Pt(0)
        date_p.paragraph_format.space_after = Pt(0)
        date_p.paragraph_format.line_spacing = 1.0
        date_label_run = date_p.add_run('Generated on: ')
        date_label_run.bold = True
        date_label_run.font.color.rgb = RGBColor(*self.hex_to_rgb(self.text_color)) if not self.dark_theme else RGBColor(255, 255, 255)
        date_run = date_p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        date_run.font.color.rgb = RGBColor(*self.hex_to_rgb(self.text_color)) if not self.dark_theme else RGBColor(255, 255, 255)
        text_rgb = self.hex_to_rgb(self.text_color)
        label_rgb = self.hex_to_rgb(self.label_color)
        temp_dir = images_folder if images_folder else tempfile.mkdtemp()
        try:
            for page_start in range(0, len(subtitles), self.subtitles_per_page):
                if not self.is_processing:
                    break
                if page_start > 0 and self.page_organization:
                    doc.add_page_break()
                page_subtitles = subtitles[page_start:page_start + self.subtitles_per_page]
                page_screenshots = screenshots[page_start:page_start + self.subtitles_per_page]
                for i, (subtitle, screenshot) in enumerate(zip(page_subtitles, page_screenshots)):
                    global_idx = page_start + i
                    heading = doc.add_heading(self.get_heading_text(subtitle, global_idx), level=2)
                    heading.paragraph_format.space_before = Pt(0)
                    heading.paragraph_format.space_after = Pt(0)
                    heading.paragraph_format.line_spacing = 1.0
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(*text_rgb) if not self.dark_theme else RGBColor(255, 255, 255)
                        run.font.size = Pt(14)
                    if self.add_bookmarks:
                        bookmark_name = f"{video_name}_subtitle_{subtitle['number']}"
                        bookmark_start = OxmlElement("w:bookmarkStart")
                        bookmark_start.set(qn("w:id"), str(global_idx))
                        bookmark_start.set(qn("w:name"), bookmark_name)
                        bookmark_end = OxmlElement("w:bookmarkEnd")
                        bookmark_end.set(qn("w:id"), str(global_idx))
                        heading._p.append(bookmark_start)
                        heading._p.append(bookmark_end)
                    timing_p = doc.add_paragraph()
                    timing_p.paragraph_format.space_before = Pt(0)
                    timing_p.paragraph_format.space_after = Pt(0)
                    timing_p.paragraph_format.line_spacing = 1.0
                    timing_label_run = timing_p.add_run('Time: ')
                    timing_label_run.bold = True
                    timing_label_run.font.color.rgb = RGBColor(*text_rgb) if not self.dark_theme else RGBColor(255, 255, 255)
                    timing_label_run.font.size = Pt(10)
                    start_formatted = self.format_time(subtitle['start_time'])
                    end_formatted = self.format_time(subtitle['end_time'])
                    time_run = timing_p.add_run(f"{start_formatted} → {end_formatted}")
                    time_run.font.color.rgb = RGBColor(*text_rgb) if not self.dark_theme else RGBColor(255, 255, 255)
                    time_run.font.size = Pt(10)
                    text_p = doc.add_paragraph()
                    text_p.paragraph_format.space_before = Pt(0)
                    text_p.paragraph_format.space_after = Pt(0)
                    text_p.paragraph_format.line_spacing = 1.0
                    text_label_run = text_p.add_run('Text: ')
                    text_label_run.bold = True
                    text_label_run.font.color.rgb = RGBColor(*label_rgb)
                    text_label_run.font.size = Pt(10)
                    text_content_run = text_p.add_run(subtitle['text'])
                    text_content_run.font.color.rgb = RGBColor(*text_rgb) if not self.dark_theme else RGBColor(255, 255, 255)
                    text_content_run.font.size = Pt(10)
                    if screenshot is not None:
                        screenshot_p = doc.add_paragraph()
                        screenshot_p.paragraph_format.space_before = Pt(0)
                        screenshot_p.paragraph_format.space_after = Pt(0)
                        screenshot_p.paragraph_format.line_spacing = 1.0
                        screenshot_label_run = screenshot_p.add_run('Screenshot: ')
                        screenshot_label_run.bold = True
                        screenshot_label_run.font.color.rgb = RGBColor(*text_rgb) if not self.dark_theme else RGBColor(255, 255, 255)
                        screenshot_label_run.font.size = Pt(10)
                        img_path = os.path.join(temp_dir, f'screenshot_{global_idx}.png')
                        img = Image.fromarray(screenshot)
                        img.save(img_path)
                        doc.add_picture(img_path, width=Inches(6.5))
                if not self.page_organization and page_start + self.subtitles_per_page < len(subtitles):
                    separator_p = doc.add_paragraph('─' * 50)
                    separator_p.paragraph_format.space_before = Pt(0)
                    separator_p.paragraph_format.space_after = Pt(0)
                    separator_p.paragraph_format.line_spacing = 1.0
                    for run in separator_p.runs:
                        run.font.color.rgb = RGBColor(*text_rgb) if not self.dark_theme else RGBColor(255, 255, 255)
        finally:
            if not images_folder:
                for file in os.listdir(temp_dir):
                    os.remove(os.path.join(temp_dir, file))
                os.rmdir(temp_dir)
        doc.save(output_path)
        self.log_message(f"DOCX saved to: {output_path}", log_list)

    def create_markdown_report(self, subtitles, screenshots, output_path, relative_path, base_output_dir, log_list):
        output_path, images_folder = self.create_output_folder(output_path, relative_path, base_output_dir)
        content = []
        content.append(f"# {self.custom_title or 'Video Subtitle Report'}\n")
        content.append(f"**Total subtitles:** {len(subtitles)}\n")
        content.append(f"**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        content.append(f"**Theme:** {'Dark' if self.dark_theme else 'Light'}\n")
        temp_dir = images_folder if images_folder else tempfile.mkdtemp()
        try:
            for page_start in range(0, len(subtitles), self.subtitles_per_page):
                if not self.is_processing:
                    break
                if page_start > 0 and self.page_organization:
                    content.append('<div class="page-break"></div>\n')
                page_subtitles = subtitles[page_start:page_start + self.subtitles_per_page]
                page_screenshots = screenshots[page_start:page_start + self.subtitles_per_page]
                for i, (subtitle, screenshot) in enumerate(zip(page_subtitles, page_screenshots)):
                    global_idx = page_start + i
                    heading_text = self.get_heading_text(subtitle, global_idx)
                    content.append(f"## {heading_text}\n")
                    start_formatted = self.format_time(subtitle['start_time'])
                    end_formatted = self.format_time(subtitle['end_time'])
                    content.append(f"**Time:** {start_formatted} → {end_formatted}\n")
                    content.append(f'<span class="text-label">Text: </span>')
                    content.append(f'<span class="subtitle-text">{subtitle["text"]}</span>\n')
                    if screenshot is not None:
                        img_filename = f'screenshot_{global_idx}.png'
                        img_path = os.path.join(temp_dir, img_filename)
                        img = Image.fromarray(screenshot)
                        img.save(img_path)
                        img_relative_path = f"images/{img_filename}" if images_folder else img_filename
                        content.append(f"**Screenshot:**\n")
                        content.append(f"![Screenshot {global_idx+1}]({img_relative_path})\n")
                if not self.page_organization:
                    content.append("---\n")
        finally:
            if not images_folder:
                for file in os.listdir(temp_dir):
                    temp_file_path = os.path.join(temp_dir, file)
                    if os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                os.rmdir(temp_dir)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(''.join(content))
        self.log_message(f"Markdown saved to: {output_path}", log_list)
        if images_folder:
            self.log_message(f"Images saved to: {images_folder}", log_list)

    def create_html_report(self, subtitles, screenshots, output_path, relative_path, base_output_dir, log_list):
        output_path, images_folder = self.create_output_folder(output_path, relative_path, base_output_dir)
        spacing_style = "margin: 0; padding: 0;" if self.no_spacing else "margin: 0.5em 0;"
        border_style = "margin: 10px;" if self.narrow_borders else "margin: 20px;"
        bg_color = "#1a1a1a" if self.dark_theme else "#ffffff"
        text_color = self.text_color if not self.dark_theme else "#ffffff"
        heading_color = self.text_color if not self.dark_theme else "#ffffff"
        border_color = "#333333" if self.dark_theme else "#ddd"
        highlight_bg = "#2a2a2a" if self.dark_theme else "#f9f9f9"
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.custom_title or 'Video Subtitle Report'}</title>
    <style>
        @page {{
            size: A4;
            margin: {'10mm' if self.narrow_borders else '25mm'};
        }}
        body {{
            font-family: Arial, sans-serif;
            {border_style}
            line-height: 1.2;
            background-color: {bg_color} !important;
            color: {text_color} !important;
            margin: 0;
            padding: 10px;
        }}
        h1 {{
            text-align: center;
            color: {heading_color} !important;
            margin: 0 0 10px 0;
            font-size: 24px;
        }}
        h2 {{
            color: {heading_color} !important;
            border-bottom: 1px solid {border_color};
            padding-bottom: 2px;
            margin: 5px 0;
            font-size: 16px;
        }}
        .subtitle-block {{
            {spacing_style}
            page-break-inside: avoid;
            margin-bottom: 8px;
        }}
        .page-break {{
            page-break-before: always;
        }}
        .timing {{
            font-weight: bold;
            color: {text_color} !important;
            margin: 2px 0;
            font-size: 12px;
        }}
        .text-content {{
            background: {highlight_bg};
            padding: 6px;
            border-left: 3px solid #007acc;
            page-break-inside: avoid;
            margin: 4px 0;
            font-size: 13px;
        }}
        .text-label {{
            color: {self.label_color} !important;
            font-weight: bold;
        }}
        .subtitle-text {{
            color: {text_color} !important;
        }}
        img {{
            max-width: 100%;
            height: auto;
            border: 1px solid {border_color};
            display: block;
            margin: 4px auto;
            page-break-inside: avoid;
        }}
        .separator {{
            text-align: center;
            color: #ccc;
            {spacing_style}
        }}
        .report-info {{
            background: {highlight_bg};
            padding: 10px;
            border-radius: 5px;
            margin-bottom: 10px;
            color: {text_color} !important;
            font-size: 14px;
        }}
        nav {{
            background: {highlight_bg};
            padding: 10px;
            border-radius: 5px;
            margin-bottom: 10px;
        }}
        nav a {{
            color: #007acc !important;
        }}
    </style>
</head>
<body>
    <h1>{self.custom_title or 'Video Subtitle Report'}</h1>
    <div class="report-info">
        <p><strong>Total subtitles:</strong> {len(subtitles)}</p>
        <p><strong>Generated on:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p><strong>Theme:</strong> {'Dark' if self.dark_theme else 'Light'}</p>
    </div>
"""
        if self.add_bookmarks:
            html_content += f"\n<nav><h3>Bookmarks</h3><ul>\n"
            for i, subtitle in enumerate(subtitles):
                heading_text = self.get_heading_text(subtitle, i)
                html_content += f'<li><a href="#subtitle_{subtitle["number"]}">{heading_text}</a></li>\n'
            html_content += "</ul></nav>\n"
        temp_dir = images_folder if images_folder else tempfile.mkdtemp()
        try:
            for page_start in range(0, len(subtitles), self.subtitles_per_page):
                if not self.is_processing:
                    break
                page_subtitles = subtitles[page_start:page_start + self.subtitles_per_page]
                page_screenshots = screenshots[page_start:page_start + self.subtitles_per_page]
                if page_start > 0 and self.page_organization:
                    html_content += '<div class="page-break"></div>\n'
                for i, (subtitle, screenshot) in enumerate(zip(page_subtitles, page_screenshots)):
                    global_idx = page_start + i
                    heading_text = self.get_heading_text(subtitle, global_idx)
                    bookmark_id = f'subtitle_{subtitle["number"]}' if self.add_bookmarks else ''
                    html_content += f'\n<div class="subtitle-block">\n'
                    html_content += f'<h2 id="{bookmark_id}">{heading_text}</h2>\n'
                    start_formatted = self.format_time(subtitle['start_time'])
                    end_formatted = self.format_time(subtitle['end_time'])
                    html_content += f'<p class="timing"><strong>Time:</strong> {start_formatted} → {end_formatted}</p>\n'
                    html_content += f'<div class="text-content">'
                    html_content += f'<span class="text-label">Text: </span>'
                    html_content += f'<span class="subtitle-text">{subtitle["text"].replace(chr(10), "<br>")}</span>'
                    html_content += f'</div>\n'
                    if screenshot is not None:
                        img_filename = f'screenshot_{global_idx}.png'
                        img_path = os.path.join(temp_dir, img_filename)
                        img = Image.fromarray(screenshot)
                        img.save(img_path)
                        img_src = f"images/{img_filename}" if images_folder else img_filename
                        html_content += f'<p><strong>Screenshot:</strong></p>\n'
                        html_content += f'<img src="{img_src}" alt="Screenshot {global_idx+1}" />\n'
                    html_content += '</div>\n'
                if not self.page_organization and page_start + self.subtitles_per_page < len(subtitles):
                    html_content += '<div class="separator">─────────────────────────────────────────</div>\n'
        finally:
            if not images_folder:
                for file in os.listdir(temp_dir):
                    temp_file_path = os.path.join(temp_dir, file)
                    if os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                os.rmdir(temp_dir)
        html_content += "\n</body>\n</html>"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        self.log_message(f"HTML saved to: {output_path}", log_list)
        if images_folder:
            self.log_message(f"Images saved to: {images_folder}", log_list)

    def create_epub_report(self, subtitles, screenshots, output_path, relative_path, base_output_dir, log_list):
        output_path, images_folder = self.create_output_folder(output_path, relative_path, base_output_dir)
        book = epub.EpubBook()
        book.set_identifier(str(uuid.uuid4()))
        book.set_title(self.custom_title or 'Video Subtitle Report')
        book.set_language('en')
        book.add_author('Enhanced Video Subtitle Extractor')
        css_content = f"""
            body {{ background-color: {'#1a1a1a' if self.dark_theme else '#ffffff'} !important;
                    color: {self.text_color if not self.dark_theme else '#ffffff'} !important;
                    line-height: 1.2;
                    margin: 0;
                    padding: 10px; }}
            h1, h2 {{ color: {self.text_color if not self.dark_theme else '#ffffff'} !important;
                      margin: 5px 0; }}
            .subtitle-text {{ color: {self.text_color if not self.dark_theme else '#ffffff'} !important; }}
            .text-label {{ color: {self.label_color} !important; font-weight: bold; }}
            .timing {{ color: {self.text_color if not self.dark_theme else '#ffffff'} !important;
                      font-weight: bold;
                      margin: 2px 0;
                      font-size: 12px; }}
            .report-info {{ background: {'#2a2a2a' if self.dark_theme else '#f9f9f9'};
                           padding: 10px;
                           border-radius: 5px;
                           margin-bottom: 10px; }}
            img {{ max-width: 100%;
                   border: 1px solid {'#333' if self.dark_theme else '#ddd'};
                   margin: 4px auto; }}
            .subtitle-block {{ margin-bottom: 8px; }}
        """
        nav_css = epub.EpubItem(uid="nav_css", file_name="style/nav.css", media_type="text/css", content=css_content.encode('utf-8'))
        book.add_item(nav_css)
        intro_content = f"""
        <html><head><link rel="stylesheet" href="style/nav.css"/></head><body>
        <h1>{self.custom_title or 'Video Subtitle Report'}</h1>
        <div class="report-info">
            <p><strong>Total subtitles:</strong> {len(subtitles)}</p>
            <p><strong>Generated on:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            <p><strong>Theme:</strong> {'Dark' if self.dark_theme else 'Light'}</p>
        </div>
        </body></html>
        """
        intro_chapter = epub.EpubHtml(title='Introduction', file_name='intro.xhtml', lang='en')
        intro_chapter.content = intro_content.encode('utf-8')
        intro_chapter.add_item(nav_css)
        book.add_item(intro_chapter)
        chapters = [intro_chapter]
        temp_dir = images_folder if images_folder else tempfile.mkdtemp()
        try:
            for page_start in range(0, len(subtitles), self.subtitles_per_page):
                if not self.is_processing:
                    break
                page_subtitles = subtitles[page_start:page_start + self.subtitles_per_page]
                page_screenshots = screenshots[page_start:page_start + self.subtitles_per_page]
                chapter_content = f'<html><head><link rel="stylesheet" href="style/nav.css"/></head><body>'
                for i, (subtitle, screenshot) in enumerate(zip(page_subtitles, page_screenshots)):
                    global_idx = page_start + i
                    heading_text = self.get_heading_text(subtitle, global_idx)
                    chapter_content += f'<div class="subtitle-block">\n'
                    chapter_content += f'<h2>{heading_text}</h2>\n'
                    start_formatted = self.format_time(subtitle['start_time'])
                    end_formatted = self.format_time(subtitle['end_time'])
                    chapter_content += f'<p class="timing"><strong>Time:</strong> {start_formatted} → {end_formatted}</p>\n'
                    chapter_content += f'<p><span class="text-label">Text: </span>'
                    chapter_content += f'<span class="subtitle-text">{subtitle["text"].replace(chr(10), "<br/>")}</span></p>\n'
                    if screenshot is not None:
                        img_filename = f'screenshot_{global_idx}.png'
                        img_path = os.path.join(temp_dir, img_filename)
                        img = Image.fromarray(screenshot)
                        img.save(img_path)
                        with open(img_path, 'rb') as img_file:
                            img_data = img_file.read()
                        epub_img = epub.EpubItem(uid=f"img_{global_idx}", file_name=f"images/{img_filename}",
                                              media_type="image/png", content=img_data)
                        book.add_item(epub_img)
                        chapter_content += f'<p><strong>Screenshot:</strong></p>\n'
                        chapter_content += f'<img src="images/{img_filename}" alt="Screenshot {global_idx+1}" />\n'
                    chapter_content += '</div>\n'
                chapter_content += '</body></html>'
                chapter = epub.EpubHtml(title=f"Section_{page_start//self.subtitles_per_page + 1}",
                                     file_name=f'chapter_{page_start//self.subtitles_per_page + 1}.xhtml',
                                     lang='en')
                chapter.content = chapter_content.encode('utf-8')
                chapter.add_item(nav_css)
                book.add_item(chapter)
                chapters.append(chapter)
        finally:
            if not images_folder:
                for file in os.listdir(temp_dir):
                    temp_file_path = os.path.join(temp_dir, file)
                    if os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                os.rmdir(temp_dir)
        book.toc = chapters
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ['nav'] + chapters
        epub.write_epub(output_path, book, {})
        self.log_message(f"EPUB saved to: {output_path}", log_list)

    def create_pdf_report(self, subtitles, screenshots, output_path, relative_path, base_output_dir, log_list):
        final_output_path, images_folder = self.create_output_folder(output_path, relative_path, base_output_dir)
        temp_dir = tempfile.mkdtemp()
        try:
            html_filename = f"temp_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            html_path = os.path.join(temp_dir, html_filename)
            self.create_html_for_pdf(subtitles, screenshots, html_path, images_folder or temp_dir)
            margin = '0in' if self.no_spacing else ('0.5in' if self.narrow_borders else '1in')
            pdf_css = f"""
                @page {{
                    size: A4;
                    margin: {margin};
                }}
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.0;
                    background-color: {'#1a1a1a' if self.dark_theme else '#ffffff'} !important;
                    color: {self.text_color if not self.dark_theme else '#ffffff'} !important;
                    margin: 0;
                    padding: 10px;
                }}
                h1 {{
                    text-align: center;
                    margin: 0 0 8px 0;
                    font-size: 24pt;
                }}
                h2 {{
                    color: {self.text_color if not self.dark_theme else '#ffffff'} !important;
                    border-bottom: 1px solid {'#333' if self.dark_theme else '#ddd'};
                    padding-bottom: 2px;
                    margin: 4px 0;
                    font-size: 16pt;
                }}
                .subtitle-block {{
                    page-break-inside: avoid;
                    margin-bottom: 6px;
                    padding: 6px;
                    border: 1px solid {'#333' if self.dark_theme else '#ddd'};
                    border-radius: 4px;
                }}
                .page-break {{
                    page-break-before: always;
                }}
                .timing {{
                    font-weight: bold;
                    color: {self.text_color if not self.dark_theme else '#ffffff'} !important;
                    margin: 2px 0;
                    font-size: 12pt;
                }}
                .text-content {{
                    background: {'#2a2a2a' if self.dark_theme else '#f9f9f9'};
                    padding: 6px;
                    border-left: 2px solid #007acc;
                    page-break-inside: avoid;
                    margin: 4px 0;
                    font-size: 13pt;
                    line-height: 1.0;
                }}
                .text-label {{
                    color: {self.label_color} !important;
                    font-weight: bold;
                }}
                .subtitle-text {{
                    color: {self.text_color if not self.dark_theme else '#ffffff'} !important;
                }}
                img {{
                    max-width: 100%;
                    height: auto;
                    border: 1px solid {'#333' if self.dark_theme else '#ddd'};
                    display: block;
                    margin: 4px auto;
                    page-break-inside: avoid;
                }}
                .report-info {{
                    background: {'#2a2a2a' if self.dark_theme else '#f9f9f9'};
                    padding: 8px;
                    border-radius: 4px;
                    margin-bottom: 8px;
                    color: {self.text_color if not self.dark_theme else '#ffffff'} !important;
                    font-size: 12pt;
                }}
                .screenshot-label {{
                    font-weight: bold;
                    margin: 2px 0;
                    font-size: 12pt;
                }}
            """
            if self.no_spacing:
                pdf_css += """
                body { line-height:1.0; padding:0; margin:0; }
                h1 { margin:0 0 2px 0; }
                h2 { margin:0; padding:0; }
                .subtitle-block { margin:0; padding:0; border:none; }
                .text-content { padding:2px; margin:0 0 2px 0; line-height:1.0; }
                .timing { margin:0; }
                .report-info { padding:2px; margin:0 0 2px 0; }
                .screenshot-label { margin:0; }
                img { max-width:100%; height:auto; border:none; margin:0 auto; }
                """
            html_doc = weasyprint.HTML(filename=html_path)
            html_doc.write_pdf(final_output_path, stylesheets=[weasyprint.CSS(string=pdf_css)])
            self.log_message(f"PDF saved to: {final_output_path}", log_list)
        except Exception as e:
            self.log_message(f"PDF conversion error: {str(e)}", log_list)
            raise
        finally:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)

    def create_html_for_pdf(self, subtitles, screenshots, html_path, images_dir):
        bg_color = "#1a1a1a" if self.dark_theme else "#ffffff"
        text_color = self.text_color if not self.dark_theme else "#ffffff"
        heading_color = self.text_color if not self.dark_theme else "#ffffff"
        border_color = "#333333" if self.dark_theme else "#ddd"
        highlight_bg = "#2a2a2a" if self.dark_theme else "#f9f9f9"
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{self.custom_title or 'Video Subtitle Report'}</title>
</head>
<body>
    <h1>{self.custom_title or 'Video Subtitle Report'}</h1>
    <div class="report-info">
        <p><strong>Total subtitles:</strong> {len(subtitles)}</p>
        <p><strong>Generated on:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p><strong>Theme:</strong> {'Dark' if self.dark_theme else 'Light'}</p>
    </div>
"""
        try:
            for page_start in range(0, len(subtitles), self.subtitles_per_page):
                if not self.is_processing:
                    break
                page_subtitles = subtitles[page_start:page_start + self.subtitles_per_page]
                page_screenshots = screenshots[page_start:page_start + self.subtitles_per_page]
                if page_start > 0:
                    html_content += '<div class="page-break"></div>\n'
                for i, (subtitle, screenshot) in enumerate(zip(page_subtitles, page_screenshots)):
                    global_idx = page_start + i
                    heading_text = self.get_heading_text(subtitle, global_idx)
                    html_content += f'\n<div class="subtitle-block">\n'
                    html_content += f'<h2>{heading_text}</h2>\n'
                    start_formatted = self.format_time(subtitle['start_time'])
                    end_formatted = self.format_time(subtitle['end_time'])
                    html_content += f'<p class="timing"><strong>Time:</strong> {start_formatted} → {end_formatted}</p>\n'
                    html_content += f'<div class="text-content">'
                    html_content += f'<span class="text-label">Text: </span>'
                    html_content += f'<span class="subtitle-text">{subtitle["text"].replace(chr(10), "<br>")}</span>'
                    html_content += f'</div>\n'
                    if screenshot is not None:
                        img = Image.fromarray(screenshot)
                        buffer = io.BytesIO()
                        img.save(buffer, format="PNG")
                        img_data = base64.b64encode(buffer.getvalue()).decode('utf-8')
                        img_src = f"data:image/png;base64,{img_data}"
                        html_content += f'<p class="screenshot-label"><strong>Screenshot:</strong></p>\n'
                        html_content += f'<img src="{img_src}" alt="Screenshot {global_idx+1}" />\n'
                    html_content += '</div>\n'
        except Exception as e:
            raise
        html_content += "\n</body>\n</html>"
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

    def process_video(self, video_path, srt_path, base_output_dir, log_list):
        if not self.is_processing:
            return []
        try:
            if not os.path.exists(video_path):
                self.log_message(f"Skipping {video_path}: Video file not found", log_list)
                return []
            if not os.path.exists(srt_path):
                self.log_message(f"Skipping {srt_path}: SRT file not found", log_list)
                return []
            video_name = os.path.splitext(os.path.basename(video_path))[0]
            self.log_message(f"Processing video: {video_name}", log_list)
            self.log_message(f"Document title: {self.custom_title or 'Video Subtitle Report'}", log_list)
            subtitles = self.parse_srt_file(srt_path, log_list)
            if not subtitles:
                self.log_message(f"No subtitles found in {srt_path}", log_list)
                return []
            screenshots = []
            total_subtitles = len(subtitles)
            for i, subtitle in enumerate(subtitles):
                if not self.is_processing:
                    break
                screenshot = self.capture_screenshot(video_path, subtitle['end_time'], log_list)
                screenshots.append(screenshot)
            output_files = []
            format_funcs = {
                'docx': lambda: self.create_docx_report(subtitles, screenshots, f"{video_name}_subtitles.docx", "", video_name, base_output_dir, log_list),
                'pdf': lambda: self.create_pdf_report(subtitles, screenshots, f"{video_name}_subtitles.pdf", "", base_output_dir, log_list),
                'html': lambda: self.create_html_report(subtitles, screenshots, f"{video_name}_subtitles.html", "", base_output_dir, log_list),
                'md': lambda: self.create_markdown_report(subtitles, screenshots, f"{video_name}_subtitles.md", "", base_output_dir, log_list),
                'epub': lambda: self.create_epub_report(subtitles, screenshots, f"{video_name}_subtitles.epub", "", base_output_dir, log_list)
            }
            for format_name in format_funcs:
                if request.form.get(f'export_{format_name}'):
                    self.log_message(f"Creating {format_name.upper()} report...", log_list)
                    output_path = os.path.join(base_output_dir, f"{video_name}_subtitles.{format_name}")
                    format_funcs[format_name]()
                    output_files.append(output_path)
            return output_files
        except Exception as e:
            error_msg = f"Error processing {video_path}: {str(e)}"
            self.log_message(error_msg, log_list)
            raise

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Handle file uploads
        video_file = request.files.get('video_file')
        srt_file = request.files.get('srt_file')
        if not video_file or not srt_file:
            return "Please upload both video and SRT files", 400

        # Save uploaded files
        video_filename = secure_filename(video_file.filename)
        srt_filename = secure_filename(srt_file.filename)
        video_path = os.path.join(app.config['UPLOAD_FOLDER'], video_filename)
        srt_path = os.path.join(app.config['UPLOAD_FOLDER'], srt_filename)
        video_file.save(video_path)
        srt_file.save(srt_path)

        # Get options from form
        subtitles_per_page = int(request.form.get('subtitles_per_page', 3))
        no_spacing = 'no_spacing' in request.form
        narrow_borders = 'narrow_borders' in request.form
        add_bookmarks = 'add_bookmarks' in request.form
        heading_style = request.form.get('heading_style', 'numbered')
        dark_theme = 'dark_theme' in request.form
        text_color = request.form.get('text_color', '#000000')
        label_color = request.form.get('label_color', '#FF0000')
        page_organization = 'page_organization' in request.form
        create_folder = 'create_folder' in request.form
        use_video_name = 'use_video_name' in request.form
        custom_title = request.form.get('custom_title', 'Video Subtitle Report')
        if use_video_name:
            custom_title = os.path.splitext(video_filename)[0]

        # Create processor
        processor = SubtitleProcessor(
            subtitles_per_page=subtitles_per_page,
            no_spacing=no_spacing,
            narrow_borders=narrow_borders,
            add_bookmarks=add_bookmarks,
            heading_style=heading_style,
            dark_theme=dark_theme,
            text_color=text_color,
            label_color=label_color,
            page_organization=page_organization,
            create_folder=create_folder,
            custom_title=custom_title
        )

        # Temp output dir
        output_dir = tempfile.mkdtemp()
        processor.is_processing = True
        log_list = []

        try:
            output_files = processor.process_video(video_path, srt_path, output_dir, log_list)
            if output_files:
                # Create zip
                zip_path = os.path.join(output_dir, 'reports.zip')
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    for file_path in output_files:
                        arcname = os.path.basename(file_path)
                        zipf.write(file_path, arcname)
                # Clean up individual files
                for file_path in output_files:
                    os.remove(file_path)
                return send_file(zip_path, as_attachment=True, download_name='video_subtitle_reports.zip',
                                 attachment_filename='video_subtitle_reports.zip')
            else:
                return "No output files generated", 400
        except Exception as e:
            return f"Processing error: {str(e)}", 500
        finally:
            processor.is_processing = False
            # Clean up uploads and temp dir
            os.remove(video_path)
            os.remove(srt_path)
            shutil.rmtree(output_dir, ignore_errors=True)

    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)